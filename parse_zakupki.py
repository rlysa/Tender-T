import time
import requests
from bs4 import BeautifulSoup
from datetime import date, datetime
import tiktoken
import mimetypes
import os
import urllib.parse

from config import *
from prompts import *


def get_text_from_file(file_name='ПРАЙС РДК ИЮЛЬ ОПТ СО СКИДКОЙ.pdf', words=None):
    try:
        extension = file_name.split('.')[-1].lower()
        file_text = ''
        if extension == 'pdf':
            import PyPDF2

            with open(file_name, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(reader.pages):
                    file_text += page.extract_text() + '\n'


        elif extension == 'docx':
            from docx import Document

            doc = Document(file_name)
            for para in doc.paragraphs:
                file_text += para.text + '\n'
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        file_text += '\t'.join(cells)

        elif extension == 'doc':
            import textract
            file_text = textract.process(file_path).decode("utf-8", errors="ignore")

        elif extension == 'txt':
            with open(file_name) as file:
                file_text = '\n'.join([i.strip() for i in file.readlines()])

        elif extension == 'xls':
            import xlrd
            book = xlrd.open_workbook(file_name, formatting_info=True)
            sheet = book.sheet_by_index(0)

            for row_idx in range(sheet.nrows):
                values = []
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)
                    if cell.ctype == xlrd.XL_CELL_EMPTY:
                        continue

                    xf_index = sheet.cell_xf_index(row_idx, col_idx)
                    xf = book.xf_list[xf_index]
                    fmt = book.format_map[xf.format_key].format_str
                    text_val = book.formatting.format_cell(cell, xf) if hasattr(book, "formatting") else str(cell.value)
                    values.append(text_val)
                if values:
                    file_text += ' '.join(values)

        elif extension == 'xlsx':
            file_text = {} if words else ''
            from openpyxl import load_workbook
            from openpyxl.utils import get_column_letter
            from openpyxl.styles.numbers import BUILTIN_FORMATS

            wb = load_workbook(file_name, data_only=True)
            sheet = wb.active
            for row in sheet.iter_rows():
                row_text = ''
                for cell in row:
                    if cell.value is not None and cell.value != 0:
                        if isinstance(cell.value, (int, float)):
                            fmt = cell.number_format
                            try:
                                row_text += str(cell.value)  if fmt == 'General' else format(cell.value, '.2f') + ' '
                            except:
                                row_text += str(cell.value) + ' '
                        else:
                            row_text += str(cell.value) + ' '
                if words:
                    for i in words:
                        if i in row_text:
                            if i not in file_text:
                                file_text[i] = []
                            file_text[i].append(row_text + '\n')
                else:
                    file_text += row_text + '\n'
        while '\n\n' in file_text:
            file_text = file_text.replace('\n\n', '\n')

        return file_text
    except Exception as e:
        print(f'Ошибка при чтении файла\n{e}')


def get_key_words(file_name='key_words.txt'):
    with open(file_name, encoding='utf-8') as file:
        key_words = [i.lower().strip() for i in file.readlines()]
    return key_words


def get_date(period=6):
    today = date.today()
    month = today.month - 1 + period
    year = today.year + month // 12
    month = month % 12 + 1
    day = min(today.day, [31, 29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][month - 1])
    end = date(year, month, day)
    return [today.strftime('%d.%m.%Y'), end.strftime('%d.%m.%Y')]


def get_url(start_date, end_date):
    morphology = ['&morphology=on', ''][0]
    page_number = ['&pageNumber=1', ''][0]  # ?????????????????????????????
    search_filter = ['&search-filter=Дате+размещения', '']  # ????????????? как будто бесполезно из-за sortBy
    sort_direction = '&sortDirection=' + ['false', 'true'][0]  # убывание/возрастание
    records_per_page = '&recordsPerPage=_' + ['10', '20', '50'][0]
    show_lots_info_hidden = '&showLotsInfoHidden=' + ['false', 'true'][0]
    sort_by = '&sortBy=' + ['UPDATE_DATE', 'PUBLISH_DATE', 'PRICE', 'RELEVANCE'][0]
    zakon = ''.join(['&fz44=on', '&fz223=on', '&ppRf615=on', ''])
    stage = ''.join(['&af=on', '&ca=on', '&pc=on', '&pa=on',
                     ''])  # Подача заявок/Работа комиссии/Закупка завершена/Закупка отменена
    currency_id_general = '-1'  # ???????????????????????????????????????
    publish_date_from = ['&publishDateFrom=' + start_date, ''][1]
    publish_date_to = ['&publishDateTo=' + start_date, ''][1]
    appl_submission_close_date_from = ['&applSubmissionCloseDateFrom=' + start_date, ''][0]
    appl_submission_close_date_to = ['&applSubmissionCloseDateTo=' + end_date, ''][0]

    url = f'https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString={morphology}{zakon}{stage}{appl_submission_close_date_from}{appl_submission_close_date_to}'
    return url


def get_cards(words):
    word_cards = {}
    cards = {}
    start_date, end_date = get_date()
    url = get_url(start_date, end_date)
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64',
               'Authorization': f'Bearer {ZAKUPKI_TOKEN}',
               'Content-Type': 'application/json'}
    for word in words:
        word_cards[word] = []
        url_word = url.replace('searchString=&', f'searchString={word.strip()}&')
        response_word = requests.get(url_word, headers=headers)
        while response_word.status_code != 200:
            time.sleep(2)
            print(f'{datetime.now()}: Переподключение')
            response_word = requests.get(url_word, headers=headers)
        soup_word = BeautifulSoup(response_word.text, 'html.parser')
        blocks = soup_word.find_all('div', class_='row no-gutters registry-entry__form mr-0')

        for block in blocks:
            number = block.find('div', class_='registry-entry__header-mid__number').get_text().strip().replace('№ ', '')
            name = block.find('div', class_='registry-entry__body-value').get_text().strip()
            cost = block.find('div', class_='price-block__value').get_text().strip()
            link_on_docs = f'https://zakupki.gov.ru{block.find('div', class_='href d-flex').find('a').get('href')}'
            cards[number] = [name, cost, link_on_docs, [], []]
            word_cards[word].append(number)

            response_docs = requests.get(link_on_docs, headers=headers)
            while response_docs.status_code != 200:
                time.sleep(2)
                print(f'{datetime.now()}: Переподключение')
                response_docs = requests.get(link_on_docs, headers=headers)
            soup_docs = BeautifulSoup(response_docs.text, 'html.parser')
            try:
                if link_on_docs[-1] == '=':
                    print('!--------------------------', url_word, number)
                    continue
                if 'noticeInfoId' in link_on_docs:
                    docs_names_block = soup_docs.find('div', class_='row pl-3').find_all('span', class_='count')
                    docs_names = [[j.find_all('a')[1].get('href'),  j.find_all('a')[1].get_text().strip()] for j in docs_names_block]
                elif 'regNumber' in link_on_docs:
                    docs_names_block = soup_docs.find('div', class_='blockFilesTabDocs').find_all('span', class_='section__value')
                    docs_names = [[j.find('a').get('href'), j.find('a').get_text().strip()] for j in docs_names_block]
                else:
                    docs_names = []
                for doc in docs_names:
                    if 'https://zakupki.gov.ru' not in doc[0]:
                        doc[0] = 'https://zakupki.gov.ru' + doc[0]
                    cards[number][3].append(doc[1])
                    cards[number][4].append(doc[0])
            except Exception as e:
                print('!--------------------------', response_docs.status_code, url_word, link_on_docs, e)
        print(f'{datetime.now()}: Поиск карточек: {len(word_cards)}/{len(words)}')

    return [word_cards, cards]


def download_file(url, file_name,  save_path='files'):
    headers = {
        'Authorization': f'Bearer {ZAKUPKI_TOKEN}',
        'User-Agent': 'Mozilla/5.0 (compatible; FileDownloader/1.0)'
    }
    response = requests.get(url, headers=headers, stream=True)
    while response.status_code != 200:
        time.sleep(2)
        print(f'{datetime.now()}: Переподключение')
        response =  requests.get(url, headers=headers, stream=True)
    content_type = response.headers.get('Content-Type', '').lower()
    ext = mimetypes.guess_extension(content_type.split(';')[0].strip()) or ''
    filename = file_name + '.'
    if 'content-disposition' in response.headers:
        cd = response.headers['content-disposition']
        raw_filename = cd.split('filename=')[-1].split(';')[0].strip().strip('"')
        filename += urllib.parse.unquote(raw_filename).split('.')[-1]
    if save_path:
        if os.path.isdir(save_path):
            filepath = os.path.join(save_path, filename)
        else:
            filepath = save_path
    else:
        filepath = filename

    if any(ftype in content_type for ftype in ['text', 'csv', 'xml', 'json']):
        response.encoding = response.apparent_encoding or 'utf-8'
        with open(filepath, 'w', encoding=response.encoding, errors='ignore') as f:
            f.write(response.text)
    else:
        with open(filepath, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
    return filepath


def make_request_to_ai(prompt, text, model=MODEL):
    def count_tokens(t):
        if 'gpt' in model:
            enc = tiktoken.encoding_for_model(model)
            tokens = enc.encode(t)
        elif 'Qwen' in model:
            try:
                from transformers import AutoTokenizer
                tokenizer  = AutoTokenizer.from_pretrained(model)
                tokens = tokenizer.encode(t)
            except Exception as e:
                print(f'Ошибка при подсчете токенов\n{e}')
        return tokens

    tokens_full_text = len(count_tokens(text))
    tokens_prompt = len(count_tokens(prompt))
    if tokens_full_text + tokens_prompt > MAX_TOKENS:
        max_tokens_text = (MAX_TOKENS - tokens_prompt)
        count = tokens_full_text // max_tokens_text + 1
        full_text = [prompt + text[i * max_tokens_text:i * max_tokens_text + MAX_TOKENS] for i in range(0, count)]
    else:
        full_text = [prompt + text]

    try:
        answer = []
        prompt_tokens, completion_tokens = 0, 0
        for part_of_text in full_text:
            # print(f'Количество токенов: {len(count_tokens(part_of_text))}')
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64',
                'Authorization': f'Bearer {AI_API_KEY}',
                'Content-Type': 'application/json'
            }
            data = {
                'model': model,
                'messages': [{'role': 'user', 'content': part_of_text}, ]
            }

            response = requests.post(AI_URL, headers=headers, json=data)
            answer.append(response.json()['choices'][0]['message']['content'])
            prompt_tokens += response.json()['usage']['prompt_tokens']
            completion_tokens += response.json()['usage']['completion_tokens']
            print(f'{datetime.now()}: Обработка запроса: {len(answer)}/{len(full_text)}')
        return ['\n'.join(answer), prompt_tokens, completion_tokens]
    except Exception as e:
        print(f'Ошибка в отправке запроса модели\n{e}')


def save_result(file_name, *result):
    try:
        # if mode == 'wb':
        #     with open(file_name, mode='wb') as file:
        #         for i in result.iter_content(chunk_size=8192):
        #             print(i, '\n\n\n\n------------------------------------------------------\n\n\n')
        #             file.write(i)
        # else:
        with open(file_name, encoding='utf-8', mode='w') as file:
            file.write('\n'.join(result))
        print(f'Результат сохранен в файл {file_name}')
    except Exception as e:
        print(f'Ошибка при сохранении результата в файл\n{e}')


def main():
    prompt_tokens, completion_tokens = 0, 0

    # категории товаров из файла (вручную) и ключевые, синонимичные слова (через ИИ по КТ)
    file_text = '\n'.join(get_key_words())
    print(f'{datetime.now()}: Выделены категории товаров')
    key_words_answ = make_request_to_ai(prompt_get_key_words, file_text)
    prompt_tokens += key_words_answ[1]
    completion_tokens += key_words_answ[2]
    # save_result('key_words_ai.txt', key_words_answ)
    key_words_list = key_words_answ[0].split('---------------------------')
    synonym_words = [word.strip() for word in key_words_list[-1].split('\n')]
    key_words = {}
    for word in key_words_list[0].strip().split('\n'):
        kw, kt = [i.strip() for i in word.split(':')]
        if kw not in key_words:
            key_words[kw] = []
        key_words[kw].append(kt)
    print(f'{datetime.now()}: Выделены ключевые слова')

    # карточки тендеров
    key_word_cards, key_cards = get_cards([i for i in key_words])
    synonym_word_cards, synonym_cards = get_cards(synonym_words)
    print(f'{datetime.now()}: Собраны все карточки.')
    save_result('key_words_cards.txt', '\n'.join([i for i in key_word_cards]),
                                                       '\n'.join([i for i in synonym_word_cards]),
                                                       '\n'.join([key_cards[i][2] for i in key_cards]),
                                                       '\n'.join([synonym_cards[i][2] for i in synonym_cards]))

    # фильтр 1 (через ИИ по названию карточек)
    cards_info = '\n'.join([f'{i}: {key_cards[i][0]} Документы: {', '.join(key_cards[i][3])}' for i in key_cards])
    true_cards_answ = make_request_to_ai(prompt_get_cards.replace('//Заменить1//', key_words_answ[0].replace('\n', ', ')), cards_info)
    prompt_tokens += true_cards_answ[1]
    completion_tokens += true_cards_answ[2]
    true_cards_list = true_cards_answ[0].split('\n')
    # файлы для карточек, прошедших фильтр 1, + фильтр 2 (по файлам)
    true_cards = {}
    for card in true_cards_list:
        try:
            card_num, doc_name = [i.strip() for i in card.strip().split(':')]
            if doc_name not in key_cards[card_num][3]:
                continue
            doc_link = key_cards[card_num][4][key_cards[card_num][3].index(doc_name)]
            true_cards[card_num] = key_cards[card_num][:3] + [doc_name, doc_link]
        except Exception as e:
            print('!------------------------', card, e) # key_cards[card_num]

    true_key_word_cards = {}
    for word in key_word_cards:
        key_word_cards[word] = [card for card in key_word_cards[word] if card in true_cards]
        if key_word_cards[word]:
            true_key_word_cards[word] = key_word_cards[word]
    print(f'{datetime.now()}: Отобраны релевантные карточки. Фильтр 1')
    for card in true_cards:
        path = download_file(true_cards[card][4], card)
        print(path, card, true_cards[card][3])
        true_cards[card].append(get_text_from_file(path))
    print(f'{datetime.now()}: Скачаны файлы из карточек')
    true_cards_info = '\n'.join([f'=========================================================\n{i}: {true_cards[i][0]} Документ: {true_cards[i][-1]}' for i in true_cards])
    true_cards_answ = make_request_to_ai(prompt_get_cards_2.replace('//Заменить1//', key_words_answ[0].replace('\n', ', ')), true_cards_info)
    print(f'{datetime.now()}: Отобраны релевантные карточки. Фильтр 2')
    prompt_tokens += true_cards_answ[1]
    completion_tokens += true_cards_answ[2]
    true_cards_answ = true_cards_answ[0].split('---------------------------')
    true_cards_answ_2 = []
    for i in true_cards_answ:
        if ':' in i:
            i_spl = i.split(':')
            if i_spl[0].strip().replace('\n', '') in true_cards:
                true_cards[i_spl[0].strip().replace('\n', '')][4] = i_spl[1].strip()
                true_cards_answ_2.append(i_spl[0].strip().replace('\n', ''))
    print('---------------------------------------------------------------')
    true_key_word_cards_2 = {}
    for word in true_key_word_cards:
        true_key_word_cards[word] = [card for card in true_key_word_cards[word] if card in true_cards_answ_2]
        if true_key_word_cards[word]:
            true_key_word_cards_2[word] = key_word_cards[word]
    # парсинг файла, подсчет маржи
    file_text = get_text_from_file('Прайс ХАТБЕР 27.08.25 цены С НДС.xlsx', list(true_key_word_cards.keys()))
    for i in file_text:
        file_text[i] = '\n'.join(file_text[i])
    for word in true_key_word_cards_2:
        for card in true_key_word_cards_2[word]:
            if card in true_cards and word in file_text:
                if true_cards[card] and true_cards[card][-1]:
                    # print(0)
                    true_key_word_cards[word] = 'Информация по нашим товарам:\n' + file_text[word] + '\nИнформация по карточкам:\n' + ' '.join([card + ': ' + true_cards[card][1] + true_cards[card][-1]])
    margin_info = '\n'.join([''.join(word) + ''.join(true_key_word_cards[word]) for word in true_key_word_cards_2])
    print(margin_info)
    margin_answ = make_request_to_ai(promt_count_margin, margin_info)
    print(f'{datetime.now()}: Подсчитана маржа')
    save_result('result.txt', margin_answ[0])
    prompt_tokens += margin_answ[1]
    completion_tokens += margin_answ[2]
    print(
        f'Общая стоимость: {prompt_tokens * COST_INPUT_TOKENS * 81} + {completion_tokens * COST_OUTPUT_TOKENS * 81} = {prompt_tokens * COST_INPUT_TOKENS * 81 + completion_tokens * COST_OUTPUT_TOKENS * 81}')
    # добавить обработку синонимичных слов


if __name__ == '__main__':
    for filename in os.listdir('files'):
        file_path = os.path.join('files', filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f'Ошибка при удалении {file_path}: {e}')
    print(f'{datetime.now()}: Запуск')
    main()
