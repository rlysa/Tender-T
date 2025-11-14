import re
import time
import requests
from bs4 import BeautifulSoup
from datetime import date, datetime
import tiktoken
import mimetypes
import os
import urllib.parse

from config import *
from src.prompts import *


def get_text_from_file(file_name='ПРАЙС РДК ИЮЛЬ ОПТ СО СКИДКОЙ.pdf', words=None):
    try:
        extension = file_name.split('.')[-1].lower()
        file_text = ''
        if extension == 'pdf':
            import PyPDF2

            with open(file_name, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(reader.pages):
                    file_text += page.extract_text() + '\n'
        elif extension == 'txt':
            with open(file_name, encoding='utf-8') as file:
                file_text = '\n'.join([i.lower().strip() for i in file.readlines()])

        elif extension == 'docx':
            from docx import Document
            import warnings

            warnings.filterwarnings("ignore", category=UserWarning, message="wmf image format is not supported")
            doc = Document(file_name)
            for para in doc.paragraphs:
                file_text += para.text + '\n'
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        file_text += '\t'.join(cells)

        elif extension == 'doc':
            import textract
            file_text = textract.process(file_path).decode("utf-8", errors="ignore")

        elif extension == 'txt':
            with open(file_name) as file:
                file_text = '\n'.join([i.strip() for i in file.readlines()])

        elif extension == 'xls':
            import xlrd
            book = xlrd.open_workbook(file_name, formatting_info=True)
            sheet = book.sheet_by_index(0)

            for row_idx in range(sheet.nrows):
                values = []
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)
                    if cell.ctype == xlrd.XL_CELL_EMPTY:
                        continue

                    xf_index = sheet.cell_xf_index(row_idx, col_idx)
                    xf = book.xf_list[xf_index]
                    fmt = book.format_map[xf.format_key].format_str
                    text_val = book.formatting.format_cell(cell, xf) if hasattr(book, "formatting") else str(cell.value)
                    values.append(text_val)
                if values:
                    file_text += ' '.join(values)

        elif extension == 'xlsx':
            file_text = {} if words else ''
            from openpyxl import load_workbook
            from openpyxl.utils import get_column_letter
            from openpyxl.styles.numbers import BUILTIN_FORMATS

            wb = load_workbook(file_name, data_only=True)
            sheet = wb.active
            for row in sheet.iter_rows():
                row_text = ''
                for cell in row:
                    if cell.value is not None and cell.value != 0:
                        if isinstance(cell.value, (int, float)):
                            fmt = cell.number_format
                            try:
                                row_text += str(cell.value)  if fmt == 'General' else format(cell.value, '.2f') + ' '
                            except:
                                row_text += str(cell.value) + ' '
                        else:
                            row_text += str(cell.value).lower().replace('\n', '') + ' '
                if words:
                    if 'цена' in row_text or 'название' in row_text or 'артикул' in row_text or 'стоимость' in row_text:
                        file_text['title'] = [row_text]
                    for i in words:
                        if i in row_text:
                            if i not in file_text:
                                file_text[i] = []
                            file_text[i].append(row_text)
                else:
                    file_text += row_text + '\n'
        while '\n\n' in file_text:
            file_text = file_text.replace('\n\n', '\n')
        return file_text
    except Exception as e:
        print(f'Ошибка при чтении файла {file_name} \n{e}')


def get_date(period=6):
    today = date.today()
    month = today.month - 1 + period
    year = today.year + month // 12
    month = month % 12 + 1
    day = min(today.day, [31, 29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][month - 1])
    end = date(year, month, day)
    return [today.strftime('%d.%m.%Y'), end.strftime('%d.%m.%Y')]


def get_url(start_date, end_date):
    morphology = ['&morphology=on', ''][0]
    page_number = ['&pageNumber=1', ''][0]  # ?????????????????????????????
    search_filter = ['&search-filter=Дате+размещения', '']  # ????????????? как будто бесполезно из-за sortBy
    sort_direction = '&sortDirection=' + ['false', 'true'][0]  # убывание/возрастание
    records_per_page = '&recordsPerPage=_' + ['10', '20', '50'][2]
    show_lots_info_hidden = '&showLotsInfoHidden=' + ['false', 'true'][0]
    sort_by = '&sortBy=' + ['UPDATE_DATE', 'PUBLISH_DATE', 'PRICE', 'RELEVANCE'][0]
    zakon = ''.join(['&fz44=on', '&fz223=on', '&ppRf615=on', ''])
    stage = ''.join(['&af=on', '&ca=on', '&pc=on', '&pa=on',
                     ''])  # Подача заявок/Работа комиссии/Закупка завершена/Закупка отменена
    currency_id_general = '-1'  # ???????????????????????????????????????
    publish_date_from = ['&publishDateFrom=' + start_date, ''][1]
    publish_date_to = ['&publishDateTo=' + start_date, ''][1]
    appl_submission_close_date_from = ['&applSubmissionCloseDateFrom=' + start_date, ''][0]
    appl_submission_close_date_to = ['&applSubmissionCloseDateTo=' + end_date, ''][0]

    url = f'https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString={morphology}{zakon}{stage}{appl_submission_close_date_from}{appl_submission_close_date_to}{records_per_page}'
    return url


def get_cards(words):
    def make_request(url):
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64',
                   'Authorization': f'Bearer {ZAKUPKI_TOKEN}',
                   'Content-Type': 'application/json'}
        try:
            response = requests.get(url, headers=headers)
            while response.status_code != 200:
                time.sleep(2)
                print(f'{datetime.now()}: Переподключение')
                response = requests.get(url, headers=headers)
            soup = BeautifulSoup(response.text, 'html.parser')
        except Exception as e:
            print(response.status_code, e)
        return soup

    word_cards, cards, urls = {}, {}, []
    start_date, end_date = get_date()
    url = get_url(start_date, end_date)
    for word in words:
        word_cards[word] = []
        url_word = url.replace('searchString=&', f'searchString={word.strip()}&')
        urls.append(url_word)
        soup_word = make_request(url_word)
        total = soup_word.find('div', class_='search-results__total').get_text()
        total = int(''.join([i for i in total if i.isdigit()]))
        pages = total // 50 + 1 if total > 50 else 1
        blocks = soup_word.find_all('div', class_='row no-gutters registry-entry__form mr-0')
        if pages > 1:
            for i in range(2, pages + 1):
                soup_pages = make_request(f'{url_word}&pageNumber={i}')
                blocks += soup_pages.find_all('div', class_='row no-gutters registry-entry__form mr-0')
        for block in blocks:
            number = block.find('div', class_='registry-entry__header-mid__number').get_text().strip().replace('№ ', '')
            name = block.find('div', class_='registry-entry__body-value').get_text().strip()
            cost = block.find('div', class_='price-block__value').get_text().strip().split(' ')[0].strip()
            cost = float(re.sub(r'\xa0', '', cost).replace(',', '.'))
            link = f'https://zakupki.gov.ru{block.find('div', class_='registry-entry__header-mid__number').find('a').get('href')}'
            if link: # == 'https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=0891200000625009433':
                link_on_docs = f'https://zakupki.gov.ru{block.find('div', class_='href d-flex').find('a').get('href')}'
                soup_info = make_request(link)
                region = soup_info.find('span', class_='section__title', string='Регион')
                region = region.find_next('span', class_='section__info').get_text(strip=True) if region else ''
                info = soup_info.find('div', class_='container', id='positionKTRU')
                if not info:
                    link_on_lots = soup_info.find('div', class_='tabsNav d-flex')
                    # print(link)
                    if 'Список лотов' in link_on_lots.get_text():
                        link_on_lots =  f'https://zakupki.gov.ru{link_on_lots.find_all('a')[1].get('href')}'
                        products = []
                    else:
                        print(link)
                        soup_docs = make_request(link_on_docs)
                        try:
                            if link_on_docs[-1] == '=':
                                print('!--------------------------', url_word, number)
                                continue
                            if 'noticeInfoId' in link_on_docs:
                                docs_names_block = soup_docs.find('div', class_='row pl-3').find_all('span', class_='count')
                                docs_names = [[j.find_all('a')[1].get('href'), j.find_all('a')[1].get_text().strip()] for j in
                                              docs_names_block]
                            elif 'regNumber' in link_on_docs:
                                docs_names_block = soup_docs.find('div', class_='blockFilesTabDocs').find_all('span',
                                                                                                              class_='section__value')
                                docs_names = [[j.find('a').get('href'), j.find('a').get_text().strip()] for j in
                                              docs_names_block]
                            else:
                                docs_names = []
                            for doc in docs_names:
                                if 'https://zakupki.gov.ru' not in doc[0]:
                                    doc[0] = 'https://zakupki.gov.ru' + doc[0]
                                cards[number][3].append(doc[1])
                                cards[number][4].append(doc[0])
                        except Exception as e:
                            print('!--------------------------', url_word, link_on_docs, e)
                    continue
                else:
                    total_pages = 1
                    script = info.find('script', text=re.compile('items'))
                    if script:
                        script_text = script.string
                        match = re.search(r'items\s*:\s*(\d+)', script_text)
                        if match:
                            items_value = int(match.group(1))
                            total_pages = items_value // 10 + (0 if items_value % 10 == 0 else 1)
                    # for i in range(2, total_pages + 1):
                    #     print(f"{link}#page-{i}")
                    #     save_result(f'page_{i}.txt', f'{make_request(f"{link}#page-{i}").contents[3]}')
                    #     break
                    # break
                    # total_pages = make_request(link + '#page-20')
                    # total_pages = info.find('a', class_='page__link')
                    # # total_pages = info.find('div', id=re.compile(r'truPagingContainer\d+')).get_text()
                    # if total_pages:
                    #     print(link, total_pages)
                    #     break
                    #     total_pages = total_pages.find_all('a')
                    #     print(total_pages)
                    # else:
                    #     print(1, link)
                    table = info.find('div', id=re.compile(r'purchaseObjectTruTable\d+'))
                    if not table:
                        continue
                    products = {}
                    rows = table.find_all('tr', class_='tableBlock__row')
                    for row in rows:
                        try:
                            if 'truInfo_' in str(row.get('class', [])) or row.find('th'):
                                continue
                            chevron = row.find('span', class_='chevronRight')
                            if not chevron:
                                continue
                            cells = row.find_all(['td', 'th'])
                            if len(cells) < 7:
                                continue
                            article = cells[1].get_text().strip().split()
                            if len(article) == 2:
                                article = article[1]
                            else:
                                article = article[0] + '-00000000'
                            pr_name = cells[2].get_text().strip().split('\n')[0]
                            count = cells[-3].get_text(strip=True)
                            if count == '':
                                count = '1'
                            cost_pr = cells[-1].get_text(strip=True)
                            description = ''
                            row_id = chevron.get('onclick', '')
                            if 'showInfo' in row_id:
                                match = re.search(r"showInfo\('([^']+)'", row_id)
                                if match:
                                    hidden_section_id = match.group(1)
                                    hidden_section = info.find('tr', class_=hidden_section_id)
                                    if hidden_section:
                                        char_table = hidden_section.find('table', class_='tableBlock')
                                        if char_table:
                                            char_rows = char_table.find_all('tr', class_='tableBlock__row')
                                            for char_row in char_rows[1:]:
                                                char_cells = char_row.find_all(['td', 'th'])
                                                if len(char_cells) >= 4:
                                                    char_name = char_cells[0].get_text(strip=True)
                                                    char_value = char_cells[1].get_text(strip=True)
                                                    char_unit = char_cells[2].get_text(strip=True)
                                                    if char_name and char_value:
                                                        if char_unit:
                                                            description += f'{char_name}: {char_value} {char_unit}; '
                                                        else:
                                                            description += f'{char_name}: {char_value}; '
                            if article in products:
                                article += f'-1'
                                while article in products:
                                    article = ''.join(article.split('-')[:-1]) + '-' + str(int(article.split('-')[-1]) + 1)
                            products[article] = {'name': pr_name, 'description': description.strip(),
                                                                   'count': float(re.sub(r'\xa0', '', count).replace(',', '.')),
                                                                   'cost': float(re.sub(r'\xa0', '', cost_pr.replace(',', '.')))}
                        except Exception as e:
                            print('!----------', link, e)
                cards[number] = {'name': name, 'region': region, 'cost': cost, 'link': link, 'products': products, 'link_on_docs': link_on_docs}
                word_cards[word].append(number)
        print(f'{datetime.now()}: Поиск карточек: {len(word_cards)}/{len(words)}')

    save_result(f'results/{dir}/3_searching_links.txt', '\n'.join(urls))
    return [word_cards, cards]


def download_file(url, file_name,  save_path='../files'):
    headers = {
        'Authorization': f'Bearer {ZAKUPKI_TOKEN}',
        'User-Agent': 'Mozilla/5.0 (compatible; FileDownloader/1.0)'
    }
    response = requests.get(url, headers=headers, stream=True)
    while response.status_code != 200:
        time.sleep(2)
        print(f'{datetime.now()}: Переподключение')
        response =  requests.get(url, headers=headers, stream=True)
    content_type = response.headers.get('Content-Type', '').lower()
    ext = mimetypes.guess_extension(content_type.split(';')[0].strip()) or ''
    filename = file_name + '.'
    if 'content-disposition' in response.headers:
        cd = response.headers['content-disposition']
        raw_filename = cd.split('filename=')[-1].split(';')[0].strip().strip('"')
        filename += urllib.parse.unquote(raw_filename).split('.')[-1]
    if save_path:
        if os.path.isdir(save_path):
            filepath = os.path.join(save_path, filename)
        else:
            filepath = save_path
    else:
        filepath = filename

    if any(ftype in content_type for ftype in ['text', 'csv', 'xml', 'json']):
        response.encoding = response.apparent_encoding or 'utf-8'
        with open(filepath, 'w', encoding=response.encoding, errors='ignore') as f:
            f.write(response.text)
    else:
        with open(filepath, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
    return filepath


def make_request_to_ai(prompt, text, model=MODEL):
    def count_tokens(t):
        if 'qwen' in model.lower():
            try:
                from transformers import AutoTokenizer
                tokenizer = AutoTokenizer.from_pretrained('Qwen/Qwen2-72B-Instruct-GPTQ-Int8')
                tokens = tokenizer.encode(t)
            except Exception as e:
                print(f'Ошибка при подсчете токенов\n{e}')
        elif 'gpt' in model:
            enc = tiktoken.encoding_for_model(model.replace('openai/', ''))
            tokens = enc.encode(t)
        return tokens

    tokens_full_text = len(count_tokens(text))
    tokens_prompt = len(count_tokens(prompt))
    enc = tiktoken.encoding_for_model('gpt-4o-mini')
    tokens = enc.encode(prompt+text)
    print('GPT TOKENS IN', len(tokens))
    if tokens_full_text + tokens_prompt > MAX_TOKENS:
        max_tokens_text = (MAX_TOKENS - tokens_prompt)
        count = tokens_full_text // max_tokens_text + 1
        full_text = [prompt + text[i * max_tokens_text:i * max_tokens_text + MAX_TOKENS] for i in range(0, count)]
    else:
        full_text = [prompt + text]

    # try:
    answer = []
    prompt_tokens, completion_tokens = 0, 0
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64',
        'Authorization': f'Bearer {AI_API_KEY}',
        'Content-Type': 'application/json'
    }
    for part_of_text in full_text:
        print(part_of_text)
        data = {
            'model': model,
            'messages': [{'role': 'user', 'content': part_of_text}, ]
        }

        response = requests.post(AI_URL, headers=headers, json=data)
        while response.status_code != 200:
            time.sleep(60)
            print(f'{datetime.now()}: Переподключение')
            response = requests.post(AI_URL, headers=headers, json=data)
        answer.append(response.json()['choices'][0]['message']['content'])
        prompt_tokens += response.json()['usage']['prompt_tokens']
        completion_tokens += response.json()['usage']['completion_tokens']
        enc = tiktoken.encoding_for_model('gpt-4o-mini')
        tokens = enc.encode(answer[0])
        print('GPT TOKENS OUT', len(tokens))
        print(f'{datetime.now()}: Обработка запроса: {len(answer)}/{len(full_text)}')
        time.sleep(60)
    return ['\n'.join(answer), prompt_tokens, completion_tokens]
    # except Exception as e:
    #     print(f'Ошибка в отправке запроса модели\n{e}')
        # if len(answer) != 0:
        #     return ['\n'.join(answer), prompt_tokens, completion_tokens]
        # print(f'Wait')
        # time.sleep(120)
        # print(f'Repeat')
        # make_request_to_ai(prompt, text)


def save_result(file_name, *result):
    try:
        with open(file_name, encoding='utf-8', mode='w') as file:
            file.write('\n'.join(result))
        print(f'Результат сохранен в файл {file_name}')
    except Exception as e:
        print(f'Ошибка при сохранении результата в файл\n{e}')


def main():
    prompt_tokens, completion_tokens = 0, 0
    os.mkdir(f'results/{dir}')

    # категории товаров из файла (вручную) и ключевые (через ИИ по КТ)
    product_categories = get_text_from_file('../input/product_category.txt')
    print(f'{datetime.now()}: Выделены категории товаров')
    # key_words = make_request_to_ai(prompt_get_key_words, product_categories)
    # prompt_tokens += key_words[1]
    # completion_tokens += key_words[2]
    # category_key_words = {}
    # for pair in key_words[0].strip().split('\n'):
    #     category, word = [i.strip() for i in pair.split(':')]
    #     if category not in category_key_words:
    #         category_key_words[category] = []
    #     category_key_words[category].append(word)
    # key_words = [word for category in category_key_words for word in category_key_words[category]]
    # save_result(f'results/{dir}/1_key_words.txt', '\n'.join(key_words))
    # save_result(f'results/{dir}/2_key_word_categories.txt', '\n'.join([f"{category}: {', '.join(category_key_words[category])}" for category in category_key_words]))
    # print(f'{datetime.now()}: Выделены ключевые слова')

    # category_key_words = dict(list(category_key_words.items())[0:2])
    # key_words = list(set([word for category in category_key_words for word in category_key_words[category]]))
    category_key_words = {'тетрадь': ['тетрадь'], 'блокнот': ['блокнот']}
    key_words = ['тетрадь', 'блокнот']
    # карточки тендеров
    key_word_cards, cards = get_cards(key_words)
    print(f'{datetime.now()}: Собраны все карточки.')
    save_result(f'results/{dir}/4_cards.txt', '\n'.join([cards[card]['link'] for card in cards]))
    category_cards = {}
    true_cards = []
    for category in category_key_words:
        category_cards_all = [card for word in category_key_words[category] for card in key_word_cards[word]]
        cards_info = '\n'.join([f'{card}: {"; ".join([cards[card]["products"][pr]["name"] for pr in cards[card]["products"]])}' for card in category_cards_all if cards[card]['products'] != {}])
        true_cards_answ = make_request_to_ai(prompt_get_cards_1.replace('//Заменить1//', category), cards_info)
        prompt_tokens += true_cards_answ[1]
        completion_tokens += true_cards_answ[2]
        true_cards_answ = [card.strip() for card in true_cards_answ[0].split('\n')]
        for card in true_cards_answ:
            if card in cards:
                true_cards.append(card)
        category_cards[category] = true_cards_answ
    print(f'{datetime.now()}: Отобраны релевантные карточки. Фильтр 1')
    save_result(f'results/{dir}/5_filter_1.txt', '\n'.join([f"{category}\n{'\n'.join([cards[card]['link'] for card in category_cards[category] if card in cards])}\n" for category in category_cards]))
    save_result(f'results/{dir}/5_filter_1_not_true.txt', '\n'.join([cards[card]['link'] for card in cards if card not in true_cards]))
    # for card in cards:
    #     print(card)
    #     print(cards[card]['name'], cards[card]['region'], cards[card]['cost'], cards[card]['link'])
    #     for pr in cards[card]['products']:
    #         print('    ', pr, ':', cards[card]['products'][pr]['name'], cards[card]['products'][pr]['cost'], cards[card]['products'][pr]['count'], cards[card]['products'][pr]['description'])
    # фильтр 1 (через ИИ)
    # # файлы
    # cards_info = '\n'.join([f'{card}: {cards[card][3]}' for card in cards if not 'products' in cards[card]])
    # print(cards_info)
    # true_files = make_request_to_ai(prompt_get_files_names, cards_info)
    # prompt_tokens += true_files[1]
    # completion_tokens += true_files[2]
    # copy_dict = {}
    # for pair in true_files[0].strip().split('\n'):
    #     card, file = [i.strip() for i in pair.split(':', 1)]
    #     if card in cards:
    #         if file in cards[card][3]:
    #             copy_dict[card] = cards[card]
    #             copy_dict[card][4] = cards[card][4][cards[card][3].index(file)]
    #             copy_dict[card][3] = file
    # true_cards = copy_dict
    # save_result(f'results/{dir}/6_files_names.txt', '\n'.join([f"{category}\n{'\n'.join([cards[card][3] + '\t' + cards[card][2] for card in category_cards[category] if card in true_cards])}\n" for category in category_cards]))
    # for card in true_cards:
    #     path = download_file(true_cards[card][4], card)
    #     true_cards[card].append(get_text_from_file(path))
    # print(f'{datetime.now()}: Скачаны файлы из карточек')

    # # фильтр 2 (через ИИ по содержимому файлов)
    # cards_info = '\n'.join([f'=========================================================\n{card} документ: {true_cards[card][-1]}' for card in true_cards if true_cards[card][-1]])
    # true_cards_info = make_request_to_ai(prompt_get_cards_2.replace('//Заменить1//',product_categories.replace('\n', ', ')), cards_info)
    # prompt_tokens += true_cards_info[1]
    # completion_tokens += true_cards_info[2]
    # copy_dict = {}
    # for pair in true_cards_info[0].strip().split('\n'):
    #     if ':' in pair:
    #         card, info = [i.strip() for i in pair.split(':', 1)]
    #         if card in true_cards:
    #             copy_dict[card] = true_cards[card]
    #             copy_dict[card][-1] = info
    # true_cards = copy_dict
    # save_result(f'results/{dir}/7_filter_2.txt', '\n'.join([f"{category}\n{'\n'.join([cards[card][2] for card in category_cards[category] if card in true_cards])}\n" for category in category_cards]))
    # print(f'{datetime.now()}: Отобраны релевантные карточки. Фильтр 2')

    # парсинг файла
    # file_text = get_text_from_file('Прайс ХАТБЕР 27.08.25 цены С НДС.xlsx', product_categories.split('\n'))
    file_text = get_text_from_file('../input/Прайс ХАТБЕР 27.08.25 цены С НДС.xlsx', ['блокнот', 'тетрадь'])
    title = '\n'.join(file_text['title'])
    file_text.pop('title')
    category_products = {}
    products = {}
    for category in file_text:
        file_text[category] = '\n'.join(file_text[category][0:100])
        # print(prompt_get_key_info_our_products + title + '\n' + file_text[word])
        # print(prompt_get_key_info_our_products + title, file_text[category])
        answer = make_request_to_ai(prompt_get_key_info_our_products + title, file_text[category])
        prompt_tokens += answer[1]
        completion_tokens += answer[2]
        category_products[category] = []
#         answer = ['''30б4aгр_26762: "premium" блокнот sketchbook 30л (15л. крафт 160г +15 черный офсет 160г) 240х240мм без линовки жесткая подложка на гребне-кот-бутерброд-; 73.69
# 30б4aгр_28220: "premium" блокнот sketchbook 30л (15л. крафт 160г +15 черный офсет 160г) 240х240мм без линовки жесткая подложка на гребне-today vibes-; 147.38
# 30б4aгр_28225: "premium" блокнот sketchbook 30л (15л. крафт 160г +15 черный офсет 160г) 240х240мм без линовки жесткая подложка на гребне-воздушное путешествие-; 147.38
# 40б4aгр_28219: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-одуванчиковые коты-; 233.33
# 40б4aгр_28247: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 160г) 240х240мм без линовки твердая подложка на гребне-девушка с серёжкой-; 221.42
# 40б4aгр_28248: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-мир в твоих глазах-; 233.33
# 40б4aгр_28249: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 160г) 240х240мм без линовки твердая подложка на гребне-девушка с татуировкой дракона-; 99.00
# 40б4aгр_28250: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-твори и вытворяй-; 233.33
# 40б4aгр_28251: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 160г) 240х240мм без линовки твердая подложка на гребне-рисуй и мечтай-; 99.00
# 40б4aгр_28278: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 160г) 240х240мм без линовки твердая подложка на гребне-музыка да винчи-; 99.00
# 40б4aгр_28280: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-black and white-; 233.33
# 40б4лтaк_28089: "premium" блокнот sketchbook 40л 220х220мм без линовки 100г/кв.м отрывная склейка глянц.ламин.тиснение жесткая подложка-дыхание весны-; 79.33
# 40б4лтaк_28259: "premium" блокнот sketchbook 40л 220х220мм без линовки 100г/кв.м отрывная склейка глянц.ламин.тиснение жесткая подложка-in the dark-; 79.33
# 40б4лтaк_28088: "premium" блокнот sketchbook 40л 220х220мм без линовки 100г/кв.м отрывная склейка глянц.ламин.тиснение жесткая подложка-тихая бухта-; 79.33
# 40б4aгр_31827: "premium" блокнот sketchbook 40л (15л. крафт 160г +25л белый офсет 100г) 240х240мм без линовки жесткая подложка на гребне-увидеть париж!-; 118.92
# 40б4aгр_28493: "premium" блокнот sketchbook 40л (15л. крафт 160г +25л белый офсет 100г) 240х240мм без линовки жесткая подложка на гребне-художник по имени гав-; 118.92
# 40б4aгр_31993: "premium" блокнот sketchbook 40л (15л. крафт 160г +25л белый офсет 100г) 240х240мм без линовки жесткая подложка на гребне-нежная-; 118.92
# 40б4aгр_31766: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-skull roses-; 233.33
# 40б4aгр_32000: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-песнь глубины-; 233.33
# 40б4aгр_31999: "premium" блокнот sketchbook 40л (20л. белый 100г+ 20л.черный 200г) 240х240мм без линовки твердая подложка на гребне-прекрасная италия-; 233.33
# 40б4лтaк_32025: "premium" блокнот sketchbook 40л 220х220мм без линовки 100г/кв.м отрывная склейка глянц.ламин.тиснение жесткая подложка-узнай себя!-(lucia heffernan); 87.28
# 40б4лтaк_32029: "premium" блокнот sketchbook 40л 220х220мм без линовки 100г/кв.м отрывная склейка глянц.ламин.тиснение жесткая подложка-котик в пледике-; 79.33
# 40б4лтaк_31916: "premium" блокнот sketchbook 40л 220х220мм без линовки 100г/кв.м отрывная склейка глянц.ламин.тиснение жесткая подложка-очаровательные глазки-; 79.33
# 32б5aгр_22169: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне -beautiful-; 83.83
# 32б5aгр_22266: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-butterfly-бабочка; 75.73
# 32б5aгр_28290: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне -коты в искусстве-; 40.00
# 32б5aгр_27811: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-спасибо бро-; 83.83
# 32б5aгр_28160: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-романтика-; 83.83
# 32б5aгр_28284: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-сумасшедшие улитки-; 75.73
# 32б5aгр_28203: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-zoomer-girl-; 75.73
# 32б5aгр_28295: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-artist-; 75.73
# 32б5aгр_29848: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-живой взгляд-; 83.83
# 32б5aгр_31138: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне -потому что капибара-; 83.83
# 32б5aгр_31994: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-прекрасные мгновения-; 83.83
# 32б5aгр_31671: "premium" блокнот sketchbook 32л а5ф 160г/кв.м без линовки запечат. оборот с пошаговыми эскизами твердая подложка на гребне-кокетка-; 83.83
# 40б5aгр_16723: "premium" блокнот sketchbook 40л а5ф 100г/кв.м без линовки запечат. оборот с пошаговыми эскизами жесткая подложка на гребне-счастливые моменты-; 25.65
# 60б5aпс_22160: "premium" блокнот sketchbook 60л а5ф 100г/кв.м без линовки пластиковая обложка на пластик.спирали -летняя прогулка-; 50.00
# 60б5aпс_22399: "premium" блокнот sketchbook 60л а5ф 100г/кв.м без линовки пластиковая обложка на пластик.спирали -нежная акварель-; 84.53
# 60б5aпс_27938: "premium" блокнот sketchbook 60л а5ф 100г/кв.м без линовки пластиковая обложка на пластик.спирали -flower-; 108.55
# 80б5aгр_23830: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-happy time-; 111.11
# 80б5aгр_23822: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне -фламинго-; 100.23
# 80б5aгр_28162: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне -парижанка-; 111.11
# 80б5aгр_28070: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-lucky girls-; 70.81
# 80б5aгр_28087: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-велосипед-; 74.36
# 80б5aгр_28279: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-underline-; 40.00
# 80б5aгр_28285: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-рисуй и мечтай-; 74.36
# 80б5aгр_33801: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-в хорошей компании-; 111.11
# 80б5aгр_34850: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-модный мишка-; 111.11
# 80б5aгр_34921: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-аромат пиона-; 111.11
# 80б5aгр_34924: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне -летняя прогулка-; 111.11
# 80б5aгр_34923: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-настроение bubble tea-; 111.11
# 80б5aгр_34436: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-вороная грация-; 111.11
# 80б5aгр_34922: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-кофейный муар-; 111.11
# 80б5aгр_34925: "premium" блокнот sketchbook 80л а5ф 100г/кв.м без линовки твердая подложка на гребне-весь ее мир-; 111.11
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лa: "premium" блокнот sketchbook 40л а5ф 165х165мм без линовки 100г/кв.м на скобе обл. мел.картон мат.ламин. серия -кофейные коты-; 55.00
# 40б5лaк: "premium" блокнот sketchpad 40л а5ф без линовки 100г/кв.м отрывная склейка глянц. ламин. жесткая подложка серия -фрутоскоп-; 43.92
# 120бб4в1_27690: бизнес-блокнот 120л а4ф 5-цв. блок клетка тв.переплет глянц. ламин.-color line-; 322.28
# 160бб4в1_30386: бизнес-блокнот 160л а4ф 5-цв. блок клетка тв.переплет глянц. ламин.-office style-; 370.67
# 80бб4влв1_14358: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин, выб уф лак-россия-; 261.19
# 80бб4влв1_14359: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин, выб уф лак-carbonstyle-; 261.19
# 80бб4лофв1_23793: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин. 3d фольга-абстракция-; 281.14
# 80бб4лофв1_23794: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин. 3d фольга -притяжение-; 281.14
# 80бб4влв1_25181: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин, выб уф лак -авокадо-; 242.68
# 80бб4в1_26400: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин.-игра престолов- (game of thrones); 274.27
# 80бб4в1_26228: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин.-лазер би- ( laser b.); 274.27
# 80бб4в1_27685: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин.-black-; 249.11
# 80бб4в1_30469: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин.-марки ссср-; 249.11
# 80бб4влв1_30209: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин, выб уф лак-inside-; 261.19
# 80бб4влв1_30282: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин, выб уф лак-яркие эмоции-; 261.19
# 80бб4в1_29679: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин.-лолометр-; 249.11
# 80бб4в1_30387: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин.-office style-; 249.11
# 80бб4в1_00009: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет глянц. ламин.-офис-очки-; 249.11
# 80бб4в1_29684: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет глянц. ламин.-в мире грез-; 249.11
# 80бб4лофв1_30470: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин. 3d фольга-fashion-; 281.14
# 80бб4лофв1_30343: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин. 3d фольга-золотые склоны гор-; 281.14
# 80бб4лофлв1_28859: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин. 3d фольга 3d лак-китайский дракон-; 295.94
# 80бб4лофлв1_29947: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет мат.ламин. 3d фольга 3d лак-драгоценные фрукты-; 295.94
# 80бб4в1: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет тиснение croco обл. бумвинил metallic золото; 242.58
# 80бб4в1: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет тиснение croco обл. бумвинил metallic золото; 242.58
# 80бб4в1: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет тиснение croco обл. бумвинил metallic золото; 242.58
# 80бб4в1: бизнес-блокнот 80л а4ф 210х290мм 5-цв. блок клетка тв.переплет тиснение croco обл. бумвинил metallic золото; 242.58''']
#         answer2 = ['''092308: тетрадь 120л а4ф клетка на спирали; 178.55
# 089236: тетрадь 12л а5ф на скобе; 13.94
# 089237: тетрадь 18л а5ф на скобе; 16.76
# 089238: тетрадь 24л а5ф на скобе; 19.71
# 089239: тетрадь 40л а5ф на скобе; 30.44
# 089241: тетрадь 60л а5ф на скобе; 42.48
# 090501: тетрадь 80л а4ф на гребне; 111.23
# 089242: тетрадь 80л а5ф на скобе; 48.68
# 089604: тетрадь 96л а4ф на скобе; 104.67
# 089772: тетрадь предметная 46л а5ф на скобе; 38.43
# 089788: тетрадь предметная 48л а5ф на скобе; 37.44
# 061092: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 461.24
# 066773: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 535.00
# 075976: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 254.78
# 075977: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 255.60
# 075978: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 254.78
# 087458: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 535.00
# 087459: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 535.00
# 087460: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 535.00
# 087462: "premium" тетрадь sketchbook 32л а3ф 290х290мм 160г/кв.м без линовки на пластик.спирали; 535.00
# 075969: "premium" тетрадь sketchbook 50л а4ф без линовки 100г/кв.м перфорация на отрыв с твердой обложкой на гребне; 290.00
# 075971: "premium" тетрадь sketchbook 50л а4ф без линовки 100г/кв.м перфорация на отрыв с твердой обложкой на гребне; 290.00
# 061028: "premium" тетрадь sketchbook 60л а4ф 210х210мм 120г/кв.м без линовки с твердой обложкой; 331.80
# 075986: "premium" тетрадь sketchbook 60л а4ф 210х210мм 120г/кв.м без линовки с твердой обложкой; 337.49
# 075987: "premium" тетрадь sketchbook 60л а4ф 210х210мм 120г/кв.м без линовки с твердой обложкой; 178.20
# 075988: "premium" тетрадь sketchbook 60л а4ф 210х210мм 120г/кв.м без линовки с твердой обложкой; 178.20
# 075989: "premium" тетрадь sketchbook 60л а4ф 210х210мм 120г/кв.м без линовки с твердой обложкой; 178.20
# 087195: "premium" тетрадь sketchbook 60л а4ф 210х210мм 120г/кв.м без линовки с твердой обложкой; 331.80
# 049285: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 150.05
# 049288: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 150.05
# 049372: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 060003: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 150.05
# 066368: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 300.10
# 066867: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 077194: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 178.20
# 077358: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 319.99
# 077360: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 319.99
# 077365: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 150.05
# 077653: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 150.05
# 077657: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 300.10
# 077660: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 150.05
# 083538: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 083540: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 083544: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 087484: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 087490: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 087491: "premium" тетрадь sketchbook 80-50л а5ф 100г/кв.м без линовки с твердой обложкой; 438.02
# 087772: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 339.99
# 087773: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 339.99
# 087774: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 339.99
# 087775: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 339.99
# 087777: "premium" тетрадь sketchbook 80л а5ф 165х200мм 100г/кв.м без линовки с твердой обложкой; 339.99
# 087792: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 319.99
# 087793: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 319.99
# 087794: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 319.99
# 087795: "premium" тетрадь sketchbook 80л а5ф 170х170мм 120г/кв.м без линовки с твердой обложкой; 319.99
# 060358: "premium" тетрадь 48л а5ф 145x205мм 70г/кв.м в точечку на скобе; 28.91
# 060343: "premium" тетрадь 48л а5ф 205х145мм 70г/кв.м в точечку на скобе; 28.91
# 060346: "premium" тетрадь 48л а5ф 205х145мм 70г/кв.м в точечку на скобе; 28.91
# 084716: комплект тетрадь предметная 24л с интерактивн.справочн.инф. 10 шт; 505.59
# 090172: комплект тетрадь предметная 24л с интерактивн.справочн.инф. 10 шт; 505.59
# 091820: комплект тетрадь предметная 24л с интерактивн.справочн.инф. 10 шт; 450.00
# 084654: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 10 шт; 750.30
# 089702: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 968.40
# 089717: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 1315.08
# 089769: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 945.29
# 089828: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 10 шт; 750.30
# 089879: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 968.40
# 089888: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 852.84
# 090158: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 900.33
# 090164: комплект тетрадь предметная 48л с интерактивн.справочн.инф. 12 шт; 852.84
# 089612: тетрадь для записи иероглифов 48л а5ф оригинальный блок на скобе; 83.63
# 089613: тетрадь для записи иероглифов 48л а5ф оригинальный блок на скобе; 83.63
# 089614: тетрадь для записи иероглифов 48л а5ф оригинальный блок на скобе; 83.63
# 092146: тетрадь для записи понятий и формул 48л а5ф оригинальный блок на скобе; 83.63
# 092147: тетрадь для записи понятий и формул 48л а5ф оригинальный блок на скобе; 83.63
# 092148: тетрадь для записи понятий и формул 48л а5ф оригинальный блок на скобе; 83.63
# 092638: тетрадь для записи понятий и формул 24л а5ф тиснение оригинальный блок на скобе; 57.33
# 092639: тетрадь для записи понятий и формул 24л а5ф тиснение оригинальный блок на скобе; 57.33
# 092640: тетрадь для записи понятий и формул 24л а5ф тиснение оригинальный блок на скобе; 57.33
# 013485: тетрадь для записи слов 24л а6ф оригинальный блок на скобе; 35.68
# 089999: тетрадь для записи слов 24л а6ф оригинальный блок на скобе; 35.68
# 090000: тетрадь для записи слов 24л а6ф оригинальный блок на скобе; 35.68
# 090001: тетрадь для записи слов 24л а6ф оригинальный блок на скобе; 35.68
# 090004: тетрадь для записи слов 24л а6ф оригинальный блок узкая линия на скобе; 35.68
# 090006: тетрадь для записи слов 24л а6ф оригинальный блок узкая линия на скобе; 35.68
# 090008: тетрадь для записи слов 24л а6ф оригинальный блок узкая линия на скобе; 35.68
# 090009: тетрадь для записи слов 24л а6ф оригинальный блок узкая линия на скобе; 35.68
# 092149: тетрадь для записи терминов и определений 48л а5ф оригинальный блок на скобе; 83.63
# 092150: тетрадь для записи терминов и определений 48л а5ф оригинальный блок на скобе; 83.63
# 092151: тетрадь для записи терминов и определений 48л а5ф оригинальный блок на скобе; 83.63
# 092152: тетрадь для записи терминов и определений 48л а5ф оригинальный блок на скобе; 83.63
# 092153: тетрадь для записи терминов и определений 48л а5ф оригинальный блок на скобе; 83.63''']
#         answer = answer if category == 'блокнот' else answer2
        for product in answer[0].strip().replace('\n\n', '\n').split('\n'):
            if len(product.strip().split(':', 1)) == 2:
                article, name_cost = [i.strip() for i in product.strip().split(':', 1)]
                if len(name_cost.split(';', 1)) == 2:
                    name, cost = [i.strip() for i in name_cost.split(';', 1)]
                    category_products[category].append(article)
                    products[article] = [name, cost]
    save_result(f'results/{dir}/8_products.txt', '\n\n'.join([category + '\n' + '\n'.join([product + ': ' + products[product][0] + '; ' + products[product][1] for product in category_products[category]]) for category in category_products]))

    # category = 'тетрадь'
    margin_info = ''
    for category in category_products:
        if category in category_cards:
            answer = make_request_to_ai(promt_count_margin.replace('//Заменить//', '\n'.join([product + ': ' + products[product][0] + '; ' + products[product][1] for product in category_products[category]])),
                                        '\n'.join([card + ': ' + '; '.join([pr + " (" + cards[card]['products'][pr]["name"] + "; " + cards[card]['products'][pr]["description"] + ")" for pr in cards[card]['products']]) for card in category_cards[category]]))
            prompt_tokens += answer[1]
            completion_tokens += answer[2]
            margin_info += answer[0]
    save_result(f'results/{dir}/9_margin_info.txt', margin_info)
    print(f'{datetime.now()}: Подготовлены данные для подсчета маржи')

    # подсчет маржи
    margin = {}
    for pair in margin_info.strip().replace('\n\n', '\n').split('\n'):
        if len(pair.strip().split(':', 1)) == 2:
            card, their_our = [i.strip() for i in pair.strip().split(':', 1)]
            if ';' in their_our and len(their_our.split(';', 1)) == 2:
                their, our = [i.strip() for i in their_our.split(';', 1)]
                if our in products:
                    if card not in margin:
                        margin[card] = {}
                    margin[card][their] = our
            else:
                print('!----------', their_our)
        else:
            print('!----------', pair)

    result = ''
    for card in margin:
        count_margin, count_coverage = 0, 0
        not_coverage = []
        try:
            result += f'[{cards[card]["name"]}]({cards[card]["link"]})\nОбщая сумма закупки: {cards[card]["cost"]}\nРегион: {cards[card]["region"]}'
            for product in cards[card]['products']:
                if product in margin[card]:
                    product_margin = cards[card]['products'][product]['cost'] - cards[card]['products'][product]['count'] * float(products[margin[card][product]][1])
                    count_margin += product_margin
                    result += f'''\n  Лот {product}: {cards[card]['products'][product]['name']} ({cards[card]['products'][product]['description']})
    Цена закупки: {cards[card]['products'][product]['cost']}₽
    Цена за единицу: {round(cards[card]['products'][product]['cost'] / cards[card]['products'][product]['count'], 2)}₽
    Запрашиваемое количество: {cards[card]['products'][product]['count']} шт.,
    Продаваемый товар: {margin[card][product]} ({products[margin[card][product]][0]})
    Цена за единицу: {products[margin[card][product]][1]}₽
    Маржа: {round(product_margin, 2)}₽
    Маржа%: {round(product_margin / cards[card]['products'][product]['cost'] * 100, 2)}%'''
                    count_coverage += cards[card]['products'][product]['cost']
                else:
                    not_coverage.append(cards[card]['products'][product]['name'] + ' (' + product + ')')
            result += f'\nОбщая маржа: {round(count_margin, 2)}₽'
            result += f'\nОбщая маржа%: {round(count_margin / cards[card]["cost"] * 100, 2)}₽%'
            result += f'\nПокрываемость: {round(count_coverage / cards[card]["cost"] * 100, 2)}%'
            result += f'\nНе покрытые товары: {"; ".join(not_coverage)}\n\n' if not_coverage else '\n\n'
        except Exception as e:
            print(card, margin[card], e)
    result = result.replace('.00 ', '').replace('.0 ', ' ').replace(';)', ')')

    save_result(f'results/{dir}/10_result.txt', result)
    print(f'Общая стоимость: {prompt_tokens / 1000 * COST_INPUT_TOKENS * 81} + {completion_tokens / 1000 * COST_OUTPUT_TOKENS * 81} = {prompt_tokens / 1000 * COST_INPUT_TOKENS * 81 + completion_tokens / 1000 * COST_OUTPUT_TOKENS * 81}')


if __name__ == '__main__':
    dir = datetime.now().strftime("%d.%m.%Y_%H.%M.%S")
    for filename in os.listdir('files'):
        file_path = os.path.join('files', filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f'Ошибка при удалении {file_path}: {e}')
    print(f'{datetime.now()}: Запуск')
    main()
